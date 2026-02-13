from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io

class ProposalGenerator:
    def __init__(self):
        self.company_name = "BRASILHOSP DISTRIBUIDORA LTDA"
        self.company_cnpj = "00.000.000/0001-00"
        self.company_address = "Rua Teste, 123, Bairro, Cidade/UF"
        self.company_contact = "(00) 0000-0000 | contato@brasilhosp.com.br"

    def create_proposal(self, licitacao, items, prices):
        document = Document()
        
        # Styles
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Header
        header = document.sections[0].header
        htable = header.add_table(1, 2, width=Inches(6))
        htable.autofit = False
        htable.columns[0].width = Inches(4)
        htable.columns[1].width = Inches(2)
        
        # Logo/Name (Left)
        cell_left = htable.cell(0, 0)
        p = cell_left.paragraphs[0]
        run = p.add_run(self.company_name)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 153) # Dark Blue
        p.add_run(f"\nCNPJ: {self.company_cnpj}").font.size = Pt(9)
        p.add_run(f"\n{self.company_address}").font.size = Pt(9)
        
        # Date (Right)
        cell_right = htable.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")

        # Title
        document.add_heading(f"PROPOSTA COMERCIAL - {licitacao.numero}/{licitacao.ano}", 0)
        
        # Client Info
        p = document.add_paragraph()
        run = p.add_run(f"Ao Órgão: {licitacao.orgao_nome}")
        run.bold = True
        p.add_run(f"\nModalidade: Pregão Eletrônico")
        p.add_run(f"\nObjeto: {licitacao.titulo}")

        document.add_paragraph() # Spacer

        # Table
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Headers
        headers = ["Item", "Descrição", "Unid.", "Qtd.", "V. Unit.", "V. Total"]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Gray background
            tcPr = hdr_cells[i]._element.tcPr
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'E7E7E7')
            tcPr.append(shd)

        # Content
        total_geral = 0.0
        
        for item in items:
            unit_price = prices.get(item.id, item.valor_unitario)
            total = unit_price * item.quantidade
            total_geral += total
            
            row = table.add_row().cells
            row[0].text = str(item.numero_item)
            row[1].text = item.descricao
            row[2].text = item.unidade
            row[3].text = str(item.quantidade)
            
            row[4].text = f"R$ {unit_price:,.2f}"
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row[5].text = f"R$ {total:,.2f}"
            row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Total Row
        row = table.add_row().cells
        row[4].text = "TOTAL GERAL:"
        row[4].paragraphs[0].runs[0].bold = True
        
        row[5].text = f"R$ {total_geral:,.2f}"
        row[5].paragraphs[0].runs[0].bold = True
        row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()
        
        # Terms
        p = document.add_paragraph()
        p.add_run("Validade da Proposta: ").bold = True
        p.add_run("60 (sessenta) dias.")
        
        p = document.add_paragraph()
        p.add_run("Prazo de Entrega: ").bold = True
        p.add_run("Conforme Edital.")
        
        p = document.add_paragraph()
        p.add_run("Pagamento: ").bold = True
        p.add_run("Conforme Edital.")
        
        p = document.add_paragraph()
        p.add_run("Dados Bancarios: ").bold = True
        p.add_run("Banco X (000), Ag 0000, CC 00000-0.")

        document.add_paragraph()
        document.add_paragraph()
        
        # Signature
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("____________________________________________________\n")
        p.add_run(self.company_name).bold = True
        p.add_run(f"\n{self.company_contact}")

        # Save to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
